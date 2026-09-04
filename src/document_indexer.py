import html
import logging
import os
import re
import sys
import time
from pathlib import Path
from typing import Any, Dict, List, Set, Union

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from llama_index.embeddings.google_genai import GoogleGenAIEmbedding
from pypdf import PdfReader

from src.config import config
from src.retrival import get_collection
from src.text_cleaner import clean_text

logger = logging.getLogger("document_indexer")

UPLOAD_DIR = PROJECT_ROOT / "uploads"
SUPPORTED_EXTENSIONS: Set[str] = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}
MAX_FILE_SIZE_BYTES = 25 * 1024 * 1024  # 25MB limit
CHUNK_SIZE = 512
CHUNK_OVERLAP = 50
BATCH_SIZE = 16
MAX_ATTEMPTS = 5


def validate_upload(file: UploadFile) -> str:
    """Validate filename and extension."""
    if not file.filename:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file must have a valid filename."
        )

    suffix = Path(file.filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file type '{suffix}'. Supported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    return suffix


async def store_upload(file: UploadFile) -> Path:
    """Validate, store uploaded file securely, and return saved path."""
    validate_upload(file)

    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    safe_filename = Path(file.filename).name
    target_path = UPLOAD_DIR / safe_filename

    content = await file.read()
    if not content or len(content.strip()) == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file is empty (0 bytes)."
        )

    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds maximum allowed size of {MAX_FILE_SIZE_BYTES // (1024 * 1024)}MB."
        )

    target_path.write_bytes(content)
    logger.info(f"Stored uploaded file: {target_path} ({len(content)} bytes)")
    return target_path


def load_document_text(file_path: Path) -> str:
    """Extract raw text from various file formats (.txt, .md, .pdf, .docx, .html)."""
    suffix = file_path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")

    elif suffix == ".docx":
        doc = DocxDocument(file_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    elif suffix == ".pdf":
        reader = PdfReader(str(file_path))
        extracted_pages = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                extracted_pages.append(page_text)
        return "\n\n".join(extracted_pages)

    elif suffix in {".html", ".htm"}:
        raw_html = file_path.read_text(encoding="utf-8", errors="ignore")
        # Strip HTML tags while preserving text
        clean_html = re.sub(r"<style[\s\S]*?</style>", "", raw_html, flags=re.IGNORECASE)
        clean_html = re.sub(r"<script[\s\S]*?</script>", "", clean_html, flags=re.IGNORECASE)
        clean_html = re.sub(r"<[^>]+>", " ", clean_html)
        return html.unescape(clean_html)

    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def chunk_document_text(text: str, source_name: str) -> List[Any]:
    """Chunk cleaned text into sentence nodes with metadata."""
    cleaned = clean_text(text)
    if not cleaned:
        raise ValueError("Document contains no extractable or readable text.")

    document = Document(
        text=cleaned,
        metadata={"source": source_name}
    )

    splitter = SentenceSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP
    )

    nodes = splitter.get_nodes_from_documents([document])
    return nodes


def embed_and_index_chunks(nodes: List[Any], source_name: str) -> int:
    """Generate Gemini embeddings for chunks and upsert them into ChromaDB."""
    if not nodes:
        return 0

    api_key = config.OPENAI_API_KEY or config.GEMINI_API_KEY
    embed_model = GoogleGenAIEmbedding(
        model_name=config.EMBEDDING_MODEL,
        api_key=api_key
    )

    # Embed chunks in batches with retry backoff
    for start in range(0, len(nodes), BATCH_SIZE):
        batch = nodes[start:start + BATCH_SIZE]
        texts = [node.text for node in batch]

        for attempt in range(MAX_ATTEMPTS):
            try:
                embeddings = embed_model.get_text_embedding_batch(texts)
                break
            except Exception as error:
                if attempt == MAX_ATTEMPTS - 1:
                    logger.error(f"Failed to generate embeddings after {MAX_ATTEMPTS} attempts: {error}")
                    raise
                wait_seconds = 2 ** attempt
                logger.warning(f"Embedding batch failed (attempt {attempt + 1}). Retrying in {wait_seconds}s...")
                time.sleep(wait_seconds)

        for node, embedding in zip(batch, embeddings):
            node.embedding = embedding

    # Store embeddings in ChromaDB
    collection = get_collection()
    safe_stem = re.sub(r"[^a-zA-Z0-9_-]", "_", Path(source_name).name)

    chunk_ids = []
    documents = []
    embeddings = []
    metadatas = []

    for idx, node in enumerate(nodes):
        chunk_id = f"{safe_stem}:chunk_{idx}"
        chunk_ids.append(chunk_id)
        documents.append(node.text)
        embeddings.append(node.embedding)
        metadatas.append(node.metadata)

    collection.upsert(
        ids=chunk_ids,
        documents=documents,
        embeddings=embeddings,
        metadatas=metadatas
    )

    logger.info(f"Indexed {len(chunk_ids)} chunks for '{source_name}' into collection '{collection.name}'")
    return len(chunk_ids)


def process_uploaded_document(path: Union[str, Path]) -> Dict[str, Any]:
    """
    Ingest, clean, chunk, embed, and index an uploaded document:
    1. Extract raw text
    2. Clean text
    3. Split into token/sentence chunks
    4. Tag chunks with source metadata
    5. Generate embeddings with Gemini
    6. Upsert into ChromaDB
    7. Return structured summary
    """
    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    source_name = file_path.name

    # 1 & 2: Load and clean
    raw_text = load_document_text(file_path)

    # 3 & 4: Chunk and tag
    nodes = chunk_document_text(raw_text, source_name=source_name)

    # 5 & 6: Embed and index
    indexed_count = embed_and_index_chunks(nodes, source_name=source_name)

    # Relative path representation for consistent output
    rel_path = f"uploads/{source_name}"

    return {
        "document": rel_path,
        "chunks": len(nodes),
        "indexed": indexed_count
    }
